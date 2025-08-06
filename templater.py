import os
import shutil
import json
import time
from datetime import datetime
from openpyxl import Workbook, load_workbook
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
from pathlib import Path
import signal
import sys
import spacy
from playwright.sync_api import sync_playwright

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Load skill dictionary
with open("skills.json", "r") as f:
    skill_dict = json.load(f)

# === Graceful Ctrl+C handling ===
def handle_interrupt(sig, frame):
    print("\n[✗] Aborted by user.")
    sys.exit(0)
signal.signal(signal.SIGINT, handle_interrupt)

# === Base paths ===
base_path = Path(__file__).resolve().parent
template_path = base_path / "template"
jobs_path = base_path / "jobs"
excel_file = base_path / "applications.xlsx"

# === Dynamic template detection ===
template_folders = [f for f in template_path.iterdir() if f.is_dir()]
template_options = {str(i+1): f for i, f in enumerate(template_folders)}

# === Input company name ===
while True:
    company = input("Enter the company name: ").strip()
    if company:
        company = company.title()
        break
    else:
        print("[✗] Company name is required.")

# === Excel setup ===
columns = ["Company", "Role", "Status", "Link", "Date", "Date of Last Contact", "Other"]
if not excel_file.exists():
    wb = Workbook()
    ws = wb.active
    ws.append(columns)
    wb.save(excel_file)

wb = load_workbook(excel_file)
ws = wb.active

# === Check for duplicates ===
duplicate_entries = [row for row in ws.iter_rows(min_row=2, values_only=True) if row[0] and row[0].strip().lower() == company.lower()]
if duplicate_entries:
    print(f"[!] Found {len(duplicate_entries)} existing entries for '{company}':")
    for i, entry in enumerate(duplicate_entries, 1):
        print(f"  {i}. Role: {entry[1]} | Date: {entry[4]} | Link: {entry[3]}")
    confirm = input("\n[?] Add another entry for this company? (y/n): ").lower()
    if confirm != 'y':
        print("[✗] Aborted.")
        exit()

# === Get job link ===
link_url = input("Enter the job application link: ").strip()

# === Playwright scraping ===
def extract_job_info_playwright(url):
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True, slow_mo=20)
            context = browser.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120 Safari/537.36")
            page = context.new_page()
            print("[+] Opening page in browser for scraping...")
            page.goto(url, timeout=20000)
            page.wait_for_selector("body", timeout=7000)
            time.sleep(3)

            if "cloudflare" in page.title().lower():
                input("[!] Cloudflare detected. Solve CAPTCHA, then press Enter to continue...")

            title = page.title()
            body = page.inner_text("body")
            browser.close()
            return title.strip(), body.strip()
    except Exception as e:
        print(f"[!] Playwright failed: {e}")
        return "Unknown Role", ""

# === Scrape job info ===
job_title, job_text = extract_job_info_playwright(link_url)
print(f"[+] Suggested role: {job_title}")
role = input("Enter the role (or press Enter to accept): ").strip() or job_title
other = input("Other notes (optional): ").strip()
today = datetime.now().strftime("%Y-%m-%d")

# === Choose template ===
print("\nSelect a template category:")
for key, folder in template_options.items():
    print(f"{key}. {folder.name}")

template_choice = input("Enter the number corresponding to your template: ").strip()
if template_choice not in template_options:
    print("[✗] Invalid template selection.")
    exit()

selected_template_path = template_options[template_choice]
template_resume = selected_template_path / "Resume.docx"
template_cover_letter = selected_template_path / "Cover_Letter.docx"

# === Create timestamped folder ===
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
company_folder = jobs_path / company / timestamp
company_folder.mkdir(parents=True, exist_ok=True)
resume_dest = company_folder / "Resume.docx"
cover_letter_dest = company_folder / "Cover_Letter.docx"

# === Copy templates ===
try:
    shutil.copy(template_resume, resume_dest)
    shutil.copy(template_cover_letter, cover_letter_dest)
    print(f"[✓] Templates copied to: {company_folder}")
except FileNotFoundError as e:
    print(f"[ERROR] {e}")
    exit()

# === Save job description to DOCX ===
scraped_doc = company_folder / "Job_Description.docx"
doc = Document()
doc.add_heading("Job Title", level=1)
doc.add_paragraph(job_title)
doc.add_heading("Job Description", level=1)
doc.add_paragraph(job_text)

# === Extract keywords ===
def extract_keywords(text):
    doc = nlp(text)
    tokens = [token.text.lower() for token in doc if token.is_alpha and not token.is_stop]
    return list(set(tokens))

def match_skills(tokens):
    matched = {}
    for category, skills in skill_dict.items():
        matched_skills = [skill for skill in skills if skill.lower() in tokens]
        if matched_skills:
            matched[category] = matched_skills
    return matched

keywords = extract_keywords(job_text)
matched = match_skills(keywords)

if matched:
    doc.add_heading("Matched Keywords by Category", level=1)
    for cat, kws in matched.items():
        doc.add_paragraph(f"{cat}: {', '.join(kws)}")
else:
    doc.add_paragraph("No known keywords matched.")

doc.save(scraped_doc)
print(f"[✓] Job description saved to {scraped_doc}")

# === Excel entry ===
ws.append([company, role, "Applied", link_url, today, "", other])
wb.save(excel_file)
print(f"[✓] Application entry for '{company}' added to Excel.")

# === Open resume for redaction ===
print("[+] Opening resume for redaction...")
os.startfile(resume_dest)
input("[?] Press Enter when done redacting and saved the resume...")

# === Convert resume to PDF and trim ===
resume_pdf_path = resume_dest.with_suffix(".pdf")
try:
    convert(str(resume_dest), str(resume_pdf_path))
    reader = PdfReader(resume_pdf_path)
    if len(reader.pages) > 1:
        writer = PdfWriter()
        writer.add_page(reader.pages[0])
        with open(resume_pdf_path, "wb") as f_out:
            writer.write(f_out)
    print("[✓] Resume PDF ready.")
except Exception as e:
    print(f"[!] Resume PDF conversion/trimming failed: {e}")

# === Cover letter? ===
generate_cl = input("\n[?] Generate cover letter as well? (y/n): ").lower()
if generate_cl == 'y':
    print("[+] Opening cover letter for redaction...")
    os.startfile(cover_letter_dest)
    input("[?] Press Enter when done redacting and saved the cover letter...")
    cover_pdf_path = cover_letter_dest.with_suffix(".pdf")
    try:
        convert(str(cover_letter_dest), str(cover_pdf_path))
        print("[✓] Cover letter PDF ready.")
    except Exception as e:
        print(f"[!] Cover letter PDF conversion failed: {e}")
else:
    print("[✓] Skipped cover letter.")
