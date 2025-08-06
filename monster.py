import os
import shutil
import subprocess
from datetime import datetime
import requests
import json
import spacy
from bs4 import BeautifulSoup
from openpyxl import Workbook, load_workbook
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

def extract_job_info(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")

        title = soup.title.string.strip() if soup.title else "Unknown Role"
        og_title = soup.find("meta", property="og:title")
        if og_title and og_title.get("content"):
            title = og_title["content"]

        desc_tags = soup.find_all(["p", "li"])
        description_text = "\n".join(
            tag.get_text(strip=True) for tag in desc_tags if tag.get_text(strip=True)
        )

        return title.strip(), description_text.strip()
    except Exception as e:
        print(f"[!] Error extracting job info: {e}")
        return "Unknown Role", ""

def extract_keywords_spacy(text, skill_dict):
    doc = nlp(text)
    tokens = [token.text.lower() for token in doc if token.is_alpha and not token.is_stop]
    tokens_set = set(tokens)

    matched_skills = {}

    for category, keywords in skill_dict.items():
        normalized_keywords = set(k.lower() for k in keywords)
        matched = sorted(normalized_keywords & tokens_set)
        if matched:
            matched_skills[category] = matched

    return matched_skills

def main():
    base_path = os.path.abspath(os.path.dirname(__file__))
    template_path = os.path.join(base_path, "template")
    jobs_path = os.path.join(base_path, "jobs")
    excel_file = os.path.join(base_path, "applications.xlsx")

    # Load skill dictionary
    skill_dict_path = os.path.join(base_path, "skills.json")
    try:
        with open(skill_dict_path, "r") as f:
            skill_dict = json.load(f)
    except Exception as e:
        print(f"[!] Failed to load skill dictionary: {e}")
        skill_dict = {}

    # Load templates
    template_folders = sorted([
        name for name in os.listdir(template_path)
        if os.path.isdir(os.path.join(template_path, name))
    ])
    template_options = {
        str(i + 1): (folder, os.path.join(template_path, folder))
        for i, folder in enumerate(template_folders)
    }

    while True:
        company = input("Enter the company name: ").strip()
        if company:
            company = company.title()
            break
        else:
            print("[✗] Company name is required. Please enter a valid name.")

    # Excel setup
    columns = ["Company", "Role", "Status", "Link", "Date", "Date of Last Contact", "Other"]
    if not os.path.exists(excel_file):
        print("[+] Excel file does not exist. Creating...")
        wb = Workbook()
        ws = wb.active
        ws.append(columns)
        wb.save(excel_file)
        print("[✓] Excel file created.")

    wb = load_workbook(excel_file)
    ws = wb.active

    duplicate_entries = [
        row for row in ws.iter_rows(min_row=2, values_only=True)
        if row[0] and row[0].strip().lower() == company.lower()
    ]
    if duplicate_entries:
        print(f"[!] Found {len(duplicate_entries)} existing entries for '{company}':")
        for i, entry in enumerate(duplicate_entries, 1):
            print(f"  {i}. Role: {entry[1]} | Date: {entry[4]} | Link: {entry[3]}")
        confirm = input("\n[?] Do you still want to add another entry for this company? (y/n): ").lower()
        if confirm != 'y':
            print("[✗] Aborted. No changes made.")
            return

    link_url = input("Enter the job application link: ").strip()
    suggested_role, job_description = extract_job_info(link_url)

    print(f"[+] Suggested role based on the job link: {suggested_role}")
    role = input("Enter the role you are applying for (press Enter to accept suggestion): ").strip()
    if not role:
        role = suggested_role

    other = input("Other notes (optional): ").strip()
    today = datetime.now().strftime("%Y-%m-%d")
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    company_root = os.path.join(jobs_path, company)
    company_folder = os.path.join(company_root, timestamp)
    os.makedirs(company_folder, exist_ok=True)
    print(f"[✓] Created job folder: {company_folder}")

    print("\nSelect a template category:")
    for key, (name, _) in template_options.items():
        print(f"{key}. {name}")
    template_choice = input("Enter the number corresponding to your template: ").strip()
    if template_choice not in template_options:
        print("[✗] Invalid template selection.")
        return

    template_name, selected_template_path = template_options[template_choice]
    template_resume = os.path.join(selected_template_path, "Resume.docx")
    template_cover_letter = os.path.join(selected_template_path, "Cover_Letter.docx")
    resume_dest = os.path.join(company_folder, "Resume.docx")
    cover_letter_dest = os.path.join(company_folder, "Cover_Letter.docx")

    try:
        shutil.copy(template_resume, resume_dest)
        shutil.copy(template_cover_letter, cover_letter_dest)
        print(f"[✓] Templates copied from '{template_name}' to: {company_folder}")
    except FileNotFoundError as e:
        print(f"[ERROR] {e}")
        return

    matched_keywords = extract_keywords_spacy(job_description, skill_dict)

    description_path = os.path.join(company_folder, "Job_Description.docx")
    try:
        doc = Document()
        doc.add_heading("Job Title", level=1)
        doc.add_paragraph(suggested_role)

        doc.add_heading("Job Description", level=1)
        doc.add_paragraph(job_description or "No description extracted.")

        doc.add_heading("Matched Keywords by Category", level=1)
        if matched_keywords:
            for category, keywords in matched_keywords.items():
                doc.add_heading(category.replace("_", " ").title(), level=2)
                doc.add_paragraph(", ".join(keywords))
        else:
            doc.add_paragraph("No known keywords matched.")

        doc.save(description_path)
        print(f"[✓] Job description saved to: {description_path}")
    except Exception as e:
        print(f"[!] Failed to save job description: {e}")

    ws.append([company, role, "Applied", link_url, today, "", other])
    wb.save(excel_file)
    print(f"[✓] Application entry for '{company}' added to Excel.")

    print("[+] Opening resume for redaction...")
    subprocess.Popen(["start", "", resume_dest], shell=True)
    input("[?] Press Enter when you're done redacting and have saved the resume...")

    resume_pdf = resume_dest.replace(".docx", ".pdf")
    try:
        convert(resume_dest, resume_pdf)
        print("[✓] Resume successfully converted to PDF.")
        reader = PdfReader(resume_pdf)
        if len(reader.pages) > 1:
            writer = PdfWriter()
            writer.add_page(reader.pages[0])
            with open(resume_pdf, "wb") as f_out:
                writer.write(f_out)
            print("[✓] Trimmed resume to 1 page.")
    except Exception as e:
        print(f"[!] Resume PDF conversion failed: {e}")

    generate_cl = input("\n[?] Do you want to generate a cover letter as well? (y/n): ").lower()
    if generate_cl == 'y':
        print("[+] Opening cover letter for redaction...")
        subprocess.Popen(["start", "", cover_letter_dest], shell=True)
        input("[?] Press Enter when you're done redacting and have saved the cover letter...")
        cover_pdf = cover_letter_dest.replace(".docx", ".pdf")
        try:
            convert(cover_letter_dest, cover_pdf)
            print("[✓] Cover letter successfully converted to PDF.")
        except Exception as e:
            print(f"[!] Cover letter PDF conversion failed: {e}")
    else:
        print("[✓] Skipped cover letter generation.")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[✗] Aborted by user (Ctrl+C).")
